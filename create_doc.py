import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Styles
title = doc.add_heading('AgriGo - Comprehensive Project Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('1. Project Overview & Working Flow', level=1)
doc.add_paragraph("AgriGo is an integrated, real-time logistics and employment platform designed specifically for the agricultural sector. It bridges the gap between Farmers and three distinct service providers: Transport Drivers, Field Labourers, and Machinery Providers.")

doc.add_heading('General Working Flow:', level=2)
flow = doc.add_paragraph()
flow.add_run("1. Registration & Role Selection: ").bold = True
flow.add_run("Users register via Firebase Auth and select their role.\n")
flow.add_run("2. Booking Creation: ").bold = True
flow.add_run("The Farmer navigates to a booking page. For transport, the farmer enters crop details, and an ML Model predicts the best vehicle.\n")
flow.add_run("3. Dispatch & Matching: ").bold = True
flow.add_run("The system creates a REQUESTED booking in Firestore. The background dispatch services actively listen for these requests and strictly filter them.\n")
flow.add_run("4. Acceptance & Navigation: ").bold = True
flow.add_run("A provider accepts the job. The app uses Google Maps and Directions APIs to draw a real-time, dynamic route from the provider to the farmer.\n")
flow.add_run("5. Completion: ").bold = True
flow.add_run("Upon arrival, an OTP is exchanged, the job is marked COMPLETED, and earnings are recorded.")

doc.add_heading('System Architecture Flow', level=2)
doc.add_paragraph("The following diagram visualizes how the Farmer App connects through the Machine Learning API and Firebase Database to dispatch jobs directly to the Provider App.")

try:
    doc.add_picture(r'C:\Users\DELL\.gemini\antigravity\brain\1f530133-d1b4-4c18-968b-84bf1e2f9b83\agrigo_architecture_1777143728040.png', width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    doc.add_paragraph(f"[Image could not be loaded: {e}]")


doc.add_heading('2. Machine Learning (ML) Integration', level=1)
doc.add_paragraph("The project utilizes Machine Learning to automate decision-making for farmers who may not know what vehicle size is appropriate for their crop yield.")

doc.add_heading('How it works:', level=2)
ml_flow = doc.add_paragraph()
ml_flow.add_run("1. Input: ").bold = True
ml_flow.add_run("In TransportBookingActivity, the farmer selects a Crop Type (e.g., Wheat) and enters a Weight in KG (e.g., 2500).\n")
ml_flow.add_run("2. API Call: ").bold = True
ml_flow.add_run("The Android app uses the Retrofit2 library to make a POST request to an external Python-based backend hosted on Render (https://machine-learning-b5hm.onrender.com/predict).\n")
ml_flow.add_run("3. The Model: ").bold = True
ml_flow.add_run("The Python backend runs a pre-trained classification model. The model was trained on historical agricultural logistics data correlating crop density, volume, and weight to vehicle capacities.\n")
ml_flow.add_run("4. Prediction: ").bold = True
ml_flow.add_run("The model processes the JSON request and returns a predicted vehicle type.\n")
ml_flow.add_run("5. Output: ").bold = True
ml_flow.add_run("The app receives the prediction (e.g., mini_truck) and automatically pre-selects that vehicle for the farmer, streamlining the UI experience.")

doc.add_heading('3. Tools & Technologies Used', level=1)
tech = doc.add_paragraph()
tech.add_run("- Frontend Framework: ").bold = True
tech.add_run("Native Android (Java), XML for layout design, Material Design UI components.\n")
tech.add_run("- Backend & Database: ").bold = True
tech.add_run("Firebase Authentication, Cloud Firestore (A real-time NoSQL database).\n")
tech.add_run("- Mapping & Location: ").bold = True
tech.add_run("Google Maps SDK, FusedLocationProviderClient, Google Directions API.\n")
tech.add_run("- Networking: ").bold = True
tech.add_run("Retrofit2 & Gson for calling the external ML Python API.\n")

doc.add_heading('4. Team Work Division (5 Members)', level=1)
team = doc.add_paragraph()
team.add_run("1. Frontend UI/UX Developer: ").bold = True
team.add_run("Understand all XML layouts, Material Design implementations, navigation flows. Focus: res/layout/*.xml, LoginActivity, FarmerHomeActivity.\n\n")
team.add_run("2. Backend & Database Architect (Firebase): ").bold = True
team.add_run("Understand the NoSQL data structures, how collections link together, and how real-time Snapshot Listeners work. Focus: MyBookingsActivity.\n\n")
team.add_run("3. Google Maps & Location Specialist: ").bold = True
team.add_run("Understand how GPS coordinates are fetched, how Google Directions API draws polylines, and how camera animations work. Focus: TrackingActivity, LaborTrackingActivity, MapUtils.\n\n")
team.add_run("4. Dispatch & Core Logic Developer: ").bold = True
team.add_run("Master the core business logic. Understand the strict matching algorithms and background services. Focus: DriverDispatchService, DriverRequestActivity.\n\n")
team.add_run("5. Machine Learning & Networking Engineer: ").bold = True
team.add_run("Understand the Retrofit networking stack, how JSON is parsed, and theoretically explain how the Python ML model was trained and deployed. Focus: RetrofitClient, MLPredictionService.")


doc.add_heading('5. Potential Viva / Interview Questions', level=1)
q1 = doc.add_paragraph()
q1.add_run("Q1: How does the real-time map tracking work without draining the battery instantly?\n").bold = True
q1.add_run("Answer: We use the FusedLocationProviderClient with a PRIORITY_HIGH_ACCURACY request. However, we throttle the update intervals and only run location listeners while the provider is 'Online' or actively on a job. Furthermore, the route recalculation (Directions API) is only triggered if the driver deviates significantly.\n\n")
q1.add_run("Q2: Why did you choose Cloud Firestore over a traditional SQL database like MySQL?\n").bold = True
q1.add_run("Answer: Firestore is a real-time NoSQL database. For a logistics app, when a driver updates their latitude/longitude, the farmer's app needs to see that instantly. Firestore's SnapshotListeners push data to the app via WebSockets instantly, without the app having to constantly poll a SQL server.\n\n")
q1.add_run("Q3: How does the Machine Learning model predict the vehicle?\n").bold = True
q1.add_run("Answer: The Android app sends the crop type and weight to our external Python API via Retrofit. The Python backend hosts a classification model trained on logistics datasets. It classifies which standard vehicle capacity is the most efficient and cost-effective fit.\n\n")
q1.add_run("Q4: What happens if a Farmer requests an 'Auto' but there are no Autos available?\n").bold = True
q1.add_run("Answer: The system uses strict matching logic. The request sits in the transport_requests collection with a status of REQUESTED. It is broadcasted exclusively to drivers whose profile vehicleType equals 'Auto'. If none are online, the request remains pending until an Auto driver logs in.\n\n")
q1.add_run("Q5: How do you handle Android Permissions for Location?\n").bold = True
q1.add_run("Answer: Since Android 6.0, we cannot just declare permissions in the Manifest. In tracking activities, we programmatically check ContextCompat.checkSelfPermission. If it's missing, we request it. If denied, we gracefully disable the map features rather than crashing the app.")


doc.save(r'C:\Users\DELL\Documents\SATHISH\AGRIGO\AgriGo_Documentation_Final.docx')
print("Saved to AgriGo_Documentation_Final.docx")
